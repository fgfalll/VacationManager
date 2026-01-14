from datetime import datetime, date, timedelta
from typing import Tuple, Optional, List, Dict
from docx import Document
from docx.shared import Inches
import os
import sys
import pymorphy3
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import MONTHS_UK, SIGNATORIES, TEMPLATES_DIRECTORY, DEPARTMENT_NAME, FACULTY_NAME, UNIVERSITY_NAME

# Initialize MorphAnalyzer
morph = pymorphy3.MorphAnalyzer(lang='uk')

def calculate_vacation_days(start_date: date, end_date: date) -> int:
    """Calculate total vacation days including start and end dates"""
    delta = end_date - start_date
    return delta.days + 1

def get_payment_phrase(start_date: date) -> str:
    """Generate payment timing phrase based on start date"""
    day = start_date.day
    month_name = MONTHS_UK[start_date.month]

    if 1 <= day <= 15:
        return f"у першій половині {month_name}"
    else:
        return f"у другій половині {month_name}"

def get_full_name_in_genitive_case(full_name: str) -> str:
    """Converts a full Ukrainian name to the genitive case."""
    parts = full_name.split()
    genitive_parts = []
    for part in parts:
        parsed_word = morph.parse(part)[0]
        genitive_form = parsed_word.inflect({'gent'})
        if genitive_form:
            genitive_parts.append(genitive_form.word.title())
        else:
            genitive_parts.append(part)
    return " ".join(genitive_parts)

def get_position_in_genitive_case(position: str) -> str:
    """Converts a Ukrainian position to the genitive case, handling multi-word titles."""
    words = position.split(' ')
    genitive_words = []
    for word in words:
        # Preserve abbreviations
        if "." in word:
            genitive_words.append(word)
            continue
        
        p = morph.parse(word)[0]
        # Inflect only nouns and adjectives
        if 'NOUN' in p.tag or 'ADJF' in p.tag:
            genitive_form = p.inflect({'gent'})
            if genitive_form:
                # Preserve original capitalization style
                if word.islower():
                    genitive_words.append(genitive_form.word)
                elif word.istitle():
                    genitive_words.append(genitive_form.word.title())
                else: # ALL CAPS or MiXeD
                    genitive_words.append(genitive_form.word.upper())
            else:
                genitive_words.append(word)
        else:
            genitive_words.append(word)
            
    return ' '.join(genitive_words)

def get_formatted_short_name(full_name: str) -> str:
    """Formats a full name to 'First Name LAST NAME'."""
    parts = full_name.split()
    if len(parts) >= 2:
        last_name = parts[0]
        first_name = parts[1]
        return f"{first_name} {last_name.upper()}"
    return full_name

def format_vacation_description(periods: List[Dict], total_days: int) -> str:
    """
    Format vacation description for document to list individual days and ranges.
    Example: "тривалістю 10 календарних днів 1, 7, 8, 14, 15, 19 – 23 травня 2025 року"
    """
    if not periods:
        return f"терміном на {total_days} календарних днів"

    all_dates = []
    for period in periods:
        current_date = period['start_date']
        while current_date <= period['end_date']:
            all_dates.append(current_date)
            current_date += timedelta(days=1)
    
    # Sort and remove duplicates (though unique dates are guaranteed by the while loop)
    all_dates.sort()

    if not all_dates:
        return f"тривалістю {total_days} календарних днів"

    # Determine common month and year (assuming all periods are within the same month/year for simplicity here)
    # If periods span multiple months/years, a more complex logic would be needed.
    # For now, we take the month/year of the first period's start date
    common_month = all_dates[0].month
    common_year = all_dates[0].year

    formatted_days_parts = []
    i = 0
    while i < len(all_dates):
        start_range = all_dates[i]
        end_range = all_dates[i]

        # Find end of consecutive range
        j = i + 1
        while j < len(all_dates) and (all_dates[j] - all_dates[j-1]).days == 1:
            end_range = all_dates[j]
            j += 1
        
        if start_range == end_range:
            formatted_days_parts.append(str(start_range.day))
        else:
            formatted_days_parts.append(f"{start_range.day} – {end_range.day}")
        
        i = j
    
    dates_str = ", ".join(formatted_days_parts)
    month_name = MONTHS_UK[common_month]
    
    # Prepend total days
    return f"тривалістю {total_days} календарних днів {dates_str} {month_name} {common_year} року"


def generate_vacation_document_enhanced(
    request_id: int,
    staff_info: dict,
    periods: List[Dict],
    total_days: int,
    leave_type: str,
    reason_text: Optional[str] = None,
    custom_description: str = "",
    output_path: Optional[str] = None
) -> str:
    """Generate a vacation request document (.docx) with enhanced date handling"""
    template_name = "template_unpaid.docx" if leave_type == "UNPAID" else "template_paid.docx"
    project_root = os.getcwd()
    template_path = os.path.join(project_root, TEMPLATES_DIRECTORY, template_name)

    if not os.path.exists(template_path):
        create_enhanced_template(template_name)

    doc = Document(template_path)

    staff_full_name = staff_info['full_name']
    staff_genitive = get_full_name_in_genitive_case(staff_full_name)
    staff_short_name_formatted = get_formatted_short_name(staff_full_name)
    position = staff_info['position']
    
    position_genitive = get_position_in_genitive_case(position)
    position_and_department_genitive = f"{position_genitive} кафедри {DEPARTMENT_NAME}"

    employment_type = staff_info.get('employment_type', '')
    employment_type_str = f"({employment_type})" if employment_type and employment_type != 'Основне місце' else ''

    # Format vacation request text
    if periods:
        vacation_text = format_vacation_description(periods, total_days)
        # The logic here for PAID/UNPAID needs adjustment to decide whether to include vacation_type_core_str
        # For .docx, the template structure might imply how to use these.
        # This part of the code might need to be adjusted based on the specific .docx template.
        # For now, keeping it as is, but noting the potential for mismatch with the new format_vacation_description
        
        # Original logic for docx template
        if leave_type == "PAID":
            request_text = f"Прошу надати мені частину щорічної відпустки {vacation_text}."
            if len(periods) == 1:
                start_date = periods[0]['start_date']
                payment_phrase = get_payment_phrase(start_date)
                request_text += f"\n\nЗаробітну плату за час відпустки прошу виплатити {payment_phrase}."
        else: # UNPAID
            request_text = f"Прошу надати мені відпустку без збереження заробітної плати {vacation_text}"
            if reason_text:
                request_text += f" у зв'язку з {reason_text}."
            else:
                request_text += "."
    else: # Custom description
        request_text = custom_description

    # Add payment phrase placeholder for template replacement
    if periods and len(periods) == 1:
        payment_phrase = get_payment_phrase(periods[0]['start_date'])
    else:
        payment_phrase = "[виплата буде розрахована окремо]"

    replacements = {
        '{{recipient}}': SIGNATORIES['recipient'],
        '{{position_and_department_genitive}}': position_and_department_genitive,
        '{{staff_genitive}}': staff_genitive,
        '{{employment_type}}': employment_type_str,
        '{{staff_short_name_formatted}}': staff_short_name_formatted,
        '{{vacation_request_text}}': request_text,
        '{{payment_phrase}}': payment_phrase,
        '{{total_days}}': str(total_days),
        '{{director}}': SIGNATORIES['director']['name'],
        '{{quality_dept}}': SIGNATORIES['quality_dept']['name'],
    }

    is_dept_head = position.lower() in ['завідувач кафедри', 'в.о. завідувача кафедри']
    if not is_dept_head:
        replacements['{{dept_head}}'] = SIGNATORIES['dept_head']['name']
    else:
        replacements['{{dept_head}}'] = ""

    for p in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in p.text:
                # Use runs to preserve formatting
                inline = p.runs
                for i in range(len(inline)):
                    if placeholder in inline[i].text:
                        text = inline[i].text.replace(placeholder, str(value))
                        inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for placeholder, value in replacements.items():
                         if placeholder in p.text:
                            inline = p.runs
                            for i in range(len(inline)):
                                if placeholder in inline[i].text:
                                    text = inline[i].text.replace(placeholder, str(value))
                                    inline[i].text = text
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if output_path is None:
        output_filename = f"vacation_{staff_info['full_name'].replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(os.getcwd(), "generated_docs", output_filename)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

def create_enhanced_template(template_name: str):
    """Create an enhanced template with the correct structure."""
    os.makedirs(TEMPLATES_DIRECTORY, exist_ok=True)
    template_path = os.path.join(TEMPLATES_DIRECTORY, template_name)

    doc = Document()
    # Right-aligned recipient
    p = doc.add_paragraph()
    p.add_run('{{recipient}}').bold = True
    p.alignment = 2 # WD_ALIGN_PARAGRAPH.RIGHT

    # Add staff info block
    doc.add_paragraph('{{position_and_department_genitive}}').alignment = 2
    doc.add_paragraph('{{staff_genitive}}').alignment = 2
    doc.add_paragraph('{{employment_type}}').alignment = 2
    
    doc.add_paragraph() # Spacer

    # Centered title
    p = doc.add_paragraph()
    p.add_run('Заява').bold = True
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Justified body text
    p = doc.add_paragraph('{{vacation_request_text}}')
    p.alignment = 3 # WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature section
    doc.add_paragraph(f"{SIGNATORIES['director']['title']}\t\t\t\t{{{{director}}}}")
    doc.add_paragraph(f"{SIGNATORIES['quality_dept']['title']}\t\t\t\t{{{{quality_dept}}}}")
    doc.add_paragraph(f"{SIGNATORIES['dept_head']['title']}\t\t\t\t{{{{dept_head}}}}")
    
    doc.add_paragraph()
    
    # Staff signature
    p = doc.add_paragraph()
    p.add_run('\t\t\t\t\t{{staff_short_name_formatted}}')
    p.alignment = 2

    p = doc.add_paragraph()
    p.add_run('«____» ____________ 202__ р.')
    p.alignment = 2

    doc.save(template_path)

def format_display_name(staff_info: dict) -> str:
    """Format staff name for display in UI"""
    name = staff_info['full_name']
    degree = staff_info.get('academic_degree', '')
    position = staff_info.get('position', '')

    result = name
    if degree and degree != 'без ступеня':
        result += f", {degree}"
    if position:
        result += f", {position}"
    return result

def generate_vacation_document_md(
    request_id: int,
    staff_info: dict,
    periods: List[Dict],
    total_days: int,
    leave_type: str,
    reason_text: Optional[str] = None,
    custom_description: str = "",
    output_path: Optional[str] = None
) -> str:
    """Generate a vacation request document (.md) based on a template."""
    
    # Read the markdown template
    template_path = 'test_template.md'
    with open(template_path, 'r', encoding='utf-8') as f:
        template_content = f.read()

    # Prepare replacement values
    staff_full_name = staff_info['full_name']
    position = staff_info['position']

    # Generate the full vacation description using the existing robust function
    # This will now create the detailed date list/range format
    period_description = format_vacation_description(periods, total_days)
    
    # Determine the type of vacation and construct the basic descriptive string
    vacation_type_core_str = ""
    # Only include this phrase if it's explicitly paid/unpaid and not just "частину щорічної відпустки"
    if leave_type == "PAID":
        vacation_type_core_str = "зі збереженням заробітної плати"
    elif leave_type == "UNPAID":
        vacation_type_core_str = "без збереження заробітної плати"
    
    if reason_text:
        vacation_type_core_str += f" у зв'язку з {reason_text}"

    # Get employment type string
    employment_type = staff_info.get('employment_type', '')
    employment_type_str = f" за {employment_type.lower()}" if employment_type and employment_type != 'Основне місце' else ''

    # Construct payment phrase if applicable
    payment_phrase = ""
    if leave_type == "PAID" and periods and len(periods) == 1:
        start_date = periods[0]['start_date']
        payment_phrase = f". Заробітну плату за час відпустки прошу виплатити {get_payment_phrase(start_date)}."
    
    # Assemble the full description for the template
    # Example: "частину щорічної відпустки тривалістю 10 календарних днів 1, 7, 8, 14, 15, 19 – 23 травня 2025 року."
    # The "частину щорічної відпустки" part should be before vacation_type_core_str
    
    full_description_parts = ["частину щорічної відпустки"]
    if vacation_type_core_str:
        full_description_parts.append(vacation_type_core_str)
    
    full_description_parts.append(period_description)
    
    if employment_type_str:
        full_description_parts.append(employment_type_str)
    
    final_full_description = " ".join(full_description_parts) + payment_phrase + "." # Add period at the end
    
    # Parse rector's name from the recipient string for a clean header
    rector_full_string = SIGNATORIES['recipient']
    rector_name = rector_full_string.split('... ')[-1] if '... ' in rector_full_string else rector_full_string

    replacements = {
        '{university_details_genitive}': f"{UNIVERSITY_NAME}, {FACULTY_NAME}",
        '{rector_name_dative}': rector_name,
        '{applicant_position_genitive}': f"{get_position_in_genitive_case(position)} {DEPARTMENT_NAME}",
        '{applicant_full_name_genitive}': get_full_name_in_genitive_case(staff_full_name),
        '{vacation_full_description}': final_full_description,
        '{date}': datetime.now().strftime("%d.%m.%Y"),
        '{applicant_position}': position,
        '{applicant_short_name_formatted}': get_formatted_short_name(staff_full_name),
    }

    # Replace placeholders
    content = template_content
    for placeholder, value in replacements.items():
        content = content.replace(placeholder, str(value))

    # Generate output filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if output_path is None:
        output_filename = f"vacation_{staff_info['full_name'].replace(' ', '_')}_{timestamp}.md"
        output_path = os.path.join("generated_docs", output_filename)

    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Save the new markdown file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)

    return output_path