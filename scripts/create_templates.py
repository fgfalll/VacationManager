"""Script to create Word document templates for VacationManager."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create templates directory
templates_dir = Path(__file__).parent.parent / "templates"
templates_dir.mkdir(exist_ok=True)

# Common styling
def setup_document(doc):
    """Setup default document styling."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)


def create_vacation_paid_template():
    """Create template for paid vacation document."""
    doc = Document()
    setup_document(doc)

    # Header - To Rector
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Ректору Полтавського державного\nаграрного університету")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ rector_name_dav }}")
    run.bold = True

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("З А Я В А")
    run.bold = True
    run.font.size = Pt(16)

    # Content
    p = doc.add_paragraph()
    p.add_run("Прошу Вас звільнити мене від виконання посадових обов'язків ")
    p.add_run("у період з ").bold = True
    p.add_run("{{ date_start }} ")
    p.add_run("по ").bold = True
    p.add_run("{{ date_end }} ")
    p.add_run("включно ({{ days_count }} днів) та надати мені відпустку ")
    p.add_run("з наступною оплатою {{ payment_period }}.").bold = True

    # Basis text
    p = doc.add_paragraph()
    p.add_run("{{ basis_text }}")

    # Custom text (optional)
    p = doc.add_paragraph()
    p.add_run("{% if custom_text %}{{ custom_text }}{% endif %}")

    # Date placeholder
    p = doc.add_paragraph()
    p.add_run("«____» ____________ 202__ р.")

    # Signature block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if hasattr(doc, 'add_run'):
        run = p.add_run("_____________________ / {{ applicant_name }} /")

    # Department head signature
    p = doc.add_paragraph()
    p.add_run("{% if show_dept_head_signature %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Завідувач кафедри")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_____________________ / {{ dept_head_name }} /")
    p = doc.add_paragraph()
    p.add_run("{% endif %}")

    # Approvers block
    p = doc.add_paragraph()
    p.add_run("{% for approver in approvers %}")
    p = doc.add_paragraph()
    p.add_run("{{ approver.position }}")

    p = doc.add_paragraph()
    p.add_run("_____________________ / {{ approver.name }} /")
    p = doc.add_paragraph()
    p.add_run("{% endfor %}")

    # Save
    output_path = templates_dir / "vacation_paid.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_vacation_unpaid_template():
    """Create template for unpaid vacation document."""
    doc = Document()
    setup_document(doc)

    # Header - To Rector
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Ректору Полтавського державного\nаграрного університету")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ rector_name_dav }}")
    run.bold = True

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("З А Я В А")
    run.bold = True
    run.font.size = Pt(16)

    # Content
    p = doc.add_paragraph()
    p.add_run("Прошу Вас звільнити мене від виконання посадових обов'язків ")
    p.add_run("у період з ").bold = True
    p.add_run("{{ date_start }} ")
    p.add_run("по ").bold = True
    p.add_run("{{ date_end }} ")
    p.add_run("включно ({{ days_count }} днів) без збереження заробітної плати.")

    # Basis text
    p = doc.add_paragraph()
    p.add_run("{{ basis_text }}")

    # Custom text (optional)
    p = doc.add_paragraph()
    p.add_run("{% if custom_text %}{{ custom_text }}{% endif %}")

    # Date placeholder
    p = doc.add_paragraph()
    p.add_run("«____» ____________ 202__ р.")

    # Signature block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_____________________ / {{ applicant_name }} /")

    # Department head signature
    p = doc.add_paragraph()
    p.add_run("{% if show_dept_head_signature %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Завідувач кафедри")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_____________________ / {{ dept_head_name }} /")
    p = doc.add_paragraph()
    p.add_run("{% endif %}")

    # Approvers block
    p = doc.add_paragraph()
    p.add_run("{% for approver in approvers %}")
    p = doc.add_paragraph()
    p.add_run("{{ approver.position }}")

    p = doc.add_paragraph()
    p.add_run("_____________________ / {{ approver.name }} /")
    p = doc.add_paragraph()
    p.add_run("{% endfor %}")

    # Save
    output_path = templates_dir / "vacation_unpaid.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_term_extension_template():
    """Create template for term extension document."""
    doc = Document()
    setup_document(doc)

    # Header - To Rector
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Ректору Полтавського державного\nаграрного університету")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ rector_name_dav }}")
    run.bold = True

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("З А Я В А")
    run.bold = True
    run.font.size = Pt(16)

    # Content
    p = doc.add_paragraph()
    p.add_run("Прошу Вас продовжити мені термін дії контракту ")
    p.add_run("з ").bold = True
    p.add_run("{{ date_start }} ")
    p.add_run("по ").bold = True
    p.add_run("{{ date_end }} ")

    p = doc.add_paragraph()
    p.add_run("на посаді ")
    p.add_run("{{ applicant_position_gen }}").bold = True
    p.add_run(" у зв'язку з необхідністю завершення науково-дослідної роботи.")

    # Custom text (optional)
    p = doc.add_paragraph()
    p.add_run("{% if custom_text %}{{ custom_text }}{% endif %}")

    # Date placeholder
    p = doc.add_paragraph()
    p.add_run("«____» ____________ 202__ р.")

    # Signature block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_____________________ / {{ applicant_name }} /")

    # Department head signature
    p = doc.add_paragraph()
    p.add_run("{% if show_dept_head_signature %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Завідувач кафедри")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_____________________ / {{ dept_head_name }} /")
    p = doc.add_paragraph()
    p.add_run("{% endif %}")

    # Approvers block
    p = doc.add_paragraph()
    p.add_run("{% for approver in approvers %}")
    p = doc.add_paragraph()
    p.add_run("{{ approver.position }}")

    p = doc.add_paragraph()
    p.add_run("_____________________ / {{ approver.name }} /")
    p = doc.add_paragraph()
    p.add_run("{% endfor %}")

    # Save
    output_path = templates_dir / "term_extension.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    print("Creating Word document templates...")
    create_vacation_paid_template()
    create_vacation_unpaid_template()
    create_term_extension_template()
    print("All templates created successfully!")
