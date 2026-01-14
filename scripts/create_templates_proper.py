"""Скрипт для створення Word шаблонів з docxtpl Jinja2 підтримкою.

На основі реального шаблону Полтавської політехніки.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Створюємо директорію для шаблонів
templates_dir = Path(__file__).parent.parent / "templates"
templates_dir.mkdir(exist_ok=True)


def setup_document(doc):
    """Налаштовує стилі документу за замовчуванням."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    # Прибираємо відступи після параграфів
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.5


def create_vacation_paid_template():
    """Створює шаблон для оплачуваної відпустки."""
    doc = Document()
    setup_document(doc)

    # Встановлюємо поля
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(0.75)

    # Шапка - Ректору (по правому краю)
    # Рядок 1: Ректору Національного університету "Полтавська політехніка імені Юрія Кондратюка"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Ректору {{ university_name }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Рядок 2: ПІБ ректора + ступінь (дачний відмінок)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ rector_name_dav }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    # Порожній рядок
    doc.add_paragraph("")

    # Посада заявника (по правому краю)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{% if applicant_degree %}{{ applicant_degree }} {% endif %}{{ applicant_position_gen }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # ПІБ заявника (родовий відмінок)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ applicant_full_name_gen }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    # Порожній рядок
    doc.add_paragraph("")

    # Заголовок (по центру)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("З А Я В А")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

    # Порожній рядок
    doc.add_paragraph("")

    # Основний текст (з червоним рядком)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.first_line_indent = Inches(0.5)  # Червоний рядок 1.27 см

    run = p.add_run("Прошу надати мені частину щорічної відпустки тривалістю ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ days_count }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" календарних днів з ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ date_start }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" по ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ date_end }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" {{ date_end_year }} року.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Другий параграф - оплата
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.first_line_indent = Inches(0.5)
    run = p.add_run("Заробітну плату за час відпустки прошу виплатити ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ payment_period_text }}.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    # Кастомний текст (за потреби)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.first_line_indent = Inches(0.5)
    run = p.add_run("{% if custom_text %}{{ custom_text }}{% endif %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Порожній рядок
    doc.add_paragraph("")

    # Підпис заявника (по центру)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{% if applicant_degree %}{{ applicant_degree }} {% endif %}{{ applicant_position_short }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("                                          {{ applicant_full_name_nom }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    # Завідувач кафедри (опціонально)
    p = doc.add_paragraph()
    run = p.add_run("{% if show_dept_head_signature %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Підпис завідувача кафедри (по центру)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ dept_head_position }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("                                          {{ dept_head_name }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("{% endif %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Погогоджувачі (цикл)
    p = doc.add_paragraph()
    run = p.add_run("{% for approver in approvers %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_paragraph("")

    # Підпис погоджувача (по центру, багаторядковий)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ approver.position }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("                                          {{ approver.name }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("{% endfor %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Зберігаємо
    output_path = templates_dir / "vacation_paid.docx"
    doc.save(output_path)
    print(f"Створено: {output_path}")


def create_vacation_unpaid_template():
    """Створює шаблон для неоплачуваної відпустки."""
    doc = Document()
    setup_document(doc)

    # Встановлюємо поля
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(0.75)

    # Шапка - Ректору (по правому краю)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Ректору {{ university_name }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ rector_name_dav }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("")

    # Посада заявника
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{% if applicant_degree %}{{ applicant_degree }} {% endif %}{{ applicant_position_gen }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ applicant_full_name_gen }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("")

    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("З А Я В А")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph("")

    # Основний текст
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.first_line_indent = Inches(0.5)

    run = p.add_run("Прошу надати мені відпустку без збереження заробітної плати ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ days_count }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" календарних днів з ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ date_start }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" по ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ date_end }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" {{ date_end_year }} року.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Кастомний текст
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.first_line_indent = Inches(0.5)
    run = p.add_run("{% if custom_text %}{{ custom_text }}{% endif %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_paragraph("")

    # Підписи
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{% if applicant_degree %}{{ applicant_degree }} {% endif %}{{ applicant_position_short }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("                                          {{ applicant_full_name_nom }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    # Завідувач кафедри
    p = doc.add_paragraph()
    run = p.add_run("{% if show_dept_head_signature %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ dept_head_position }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("                                          {{ dept_head_name }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("{% endif %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Погогоджувачі
    p = doc.add_paragraph()
    run = p.add_run("{% for approver in approvers %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ approver.position }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("                                          {{ approver.name }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("{% endfor %}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    output_path = templates_dir / "vacation_unpaid.docx"
    doc.save(output_path)
    print(f"Створено: {output_path}")


def create_term_extension_template():
    """Створює шаблон для продовження контракту."""
    doc = Document()
    setup_document(doc)

    # Встановлюємо поля
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(0.75)

    # Шапка
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Ректору {{ university_name }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ rector_name_dav }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("")

    # Посада заявника
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{% if applicant_degree %}{{ applicant_degree }} {% endif %}{{ applicant_position_gen }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ applicant_full_name_gen }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("")

    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("З А Я В А")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph("")

    # Основний текст
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.first_line_indent = Inches(0.5)
    run = p.add_run("Прошу продовжити мені термін дії контракту ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("з ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ date_start }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" по ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ date_end }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" на посаді ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    run = p.add_run("{{ applicant_position_gen }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run(" у зв'язку з {% if custom_text %}{{ custom_text }}{% else %}необхідністю завершення науково-дослідної роботи{% endif %}.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_paragraph("")

    # Підписи
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{% if applicant_degree %}{{ applicant_degree }} {% endif %}{{ applicant_position_short }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("                                          {{ applicant_full_name_nom }}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    output_path = templates_dir / "term_extension.docx"
    doc.save(output_path)
    print(f"Створено: {output_path}")


if __name__ == "__main__":
    print("Створення Word шаблонів з docxtpl Jinja2 підтримкою...")
    print("На основі реального шаблону Полтавської політехніки\n")
    create_vacation_paid_template()
    create_vacation_unpaid_template()
    create_term_extension_template()
    print("\nВсі шаблони успішно створено!")
    print(f"Шаблони збережено в: {templates_dir}")
